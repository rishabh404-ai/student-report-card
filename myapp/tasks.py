import io
import os

import docx
from celery import shared_task
from celery.decorators import task
from celery.utils.log import get_task_logger
from django.conf import settings
from django.core.files.storage import default_storage
from django.core.mail import EmailMessage, send_mail
from django.http import HttpResponse
from django.shortcuts import render
from docx.shared import RGBColor
from rest_framework import status, viewsets
from rest_framework.decorators import api_view, renderer_classes
from rest_framework.renderers import JSONRenderer
from rest_framework.response import Response

from .emails import send_email_with_attachment
from .models import *
from .serializers import *
from celery.result import AsyncResult
from subprocess import PIPE, Popen 

logger = get_task_logger(__name__)


@shared_task
def generate_student_report():
    print('====>',generate_student_report.request.id)
    student_id = Marks.objects.values('student').distinct()
    
    for count,i in enumerate(student_id):
        doc = docx.Document()
        doc.add_heading('Student Report Card', 0)
        table = doc.add_table(rows=1, cols=2)
        row = table.rows[0].cells
        row[0].text = 'Subject'
        row[1].text = 'Marks'
        
        fetch_student_email = Student.objects.get(pk=i['student'])
        student_email = fetch_student_email.student_email           
        student_subject_and_marks = Marks.objects.filter(student=i['student']).values('students_marks','student_subject','student')
        lst_of_marks = []
        lst_of_name_of_subject_and_its_marks = []

        for count,j in enumerate(student_subject_and_marks):
            sub = Marks.objects.filter(students_marks=j['students_marks'],student=i['student']).values('student_subject')       
            for z in sub:
                student_subject_id = z['student_subject']  

            lst_of_marks.append(j['students_marks'])
            name_of_subject = Subject.objects.filter(pk=j['student_subject']).values('name_of_subject')
            student_name = Student.objects.filter(pk=j['student']).values('student_name','student_roll_no','student_class')

            for k in name_of_subject:
                lst_of_name_of_subject_and_its_marks.insert(count,(k['name_of_subject'],int(j['students_marks'])))
        
        # Calculate Percentage
        total_no_of_subjects = len(lst_of_marks) * 100
        sum_of_all_subject = sum(lst_of_marks)
        percentage = (sum_of_all_subject/total_no_of_subjects)*100
        round_percentage = round(percentage) 

        if percentage > 35:
            Student.objects.filter(pk=i['student']).update(passing_status='Pass') 
                
        else:
            Student.objects.filter(pk=i['student']).update(passing_status='Fail') 

        passing_status = Student.objects.filter(pk=i['student']).values('passing_status')
            
        for status in passing_status:
            fetch_passing_status = status['passing_status'] 
        
        
        for subject, marks in lst_of_name_of_subject_and_its_marks:
            row = table.add_row().cells
            row[0].text = str(subject)
            row[1].text = str(marks)
        
        for x in student_name:                  
                student_names = x['student_name']
                student_roll_nos = x['student_roll_no']
                student_class = x['student_class']

        doc.add_heading(f'Name: {student_names} ', 3)
        doc.add_heading(f'Email ID: {student_email} ', 3)
        doc.add_heading(f'Roll No: {student_roll_nos} ', 3)    
        doc.add_heading(f'Class: {student_class} ', 3)  
        doc.add_heading(f'Passing Status: {fetch_passing_status} ', 3)     
        doc.add_heading(f'Overall Percentage: {round_percentage}%', 3)     

        doc.save(f'Report Card {student_names}.docx')
        file = open(f'Report Card {student_names}.docx', 'rb')
        data = {'report': f'Report Card {student_names}.docx','file': file, 'email':student_email}
        send_email_with_attachment(data)
        os.remove(f'Report Card {student_names}.docx')

    print('Mail sent successfully')
    print('====>',generate_student_report.request.id)
    res = AsyncResult(generate_student_report.request.id).state
    CeleryTaskJobStatus.objects.create(job_id=generate_student_report.request.id, job_status=res)
    print('===Res===>',res)





